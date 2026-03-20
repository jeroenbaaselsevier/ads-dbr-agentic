#!/usr/bin/env python3
"""Populate Table XXX Word templates from OR table CSV outputs.

Usage example:
    .venv/bin/python projects/2026_INTERNAL_retraction/scripts/populate_or_table_docx.py \
        --template "projects/2026_INTERNAL_retraction/Table XXX.docx" \
        --csv-base projects/2026_INTERNAL_retraction/output/retraction_or_tables_paperyears_20260319 \
        --output-dir projects/2026_INTERNAL_retraction/output \
        --suffix paperyears
"""

from __future__ import annotations

import argparse
import copy
import csv
import glob
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def read_or_rows(pattern: str) -> dict[tuple[str, str], dict[str, str]]:
    rows: dict[tuple[str, str], dict[str, str]] = {}
    for file_path in sorted(glob.glob(pattern)):
        with open(file_path, encoding="utf-8") as fh:
            for row in csv.DictReader(fh):
                if row.get("variable") == "variable":
                    continue
                key = (row["variable"].strip(), row["level"].strip())
                rows[key] = row
    return rows


def format_ci(ci_str: str) -> str:
    m = re.match(r"\[([\d.]+),\s*([\d.]+)\]", ci_str.strip())
    if not m:
        return ci_str
    lower = float(m.group(1))
    upper = float(m.group(2))
    return f"{lower:.2f}-{upper:.2f}"


def format_or(or_str: str) -> str:
    if not or_str or or_str == "Ref":
        return or_str
    try:
        return f"{float(or_str):.2f}"
    except ValueError:
        return or_str


def get_caption_rprs(template_doc: Document):
    para = next(p for p in template_doc.paragraphs if p.text.strip())
    bold_rpr = None
    nonbold_rpr = None
    for r in para._element.findall(f"{{{W_NS}}}r"):
        rpr = r.find(f"{{{W_NS}}}rPr")
        if rpr is None:
            continue
        if rpr.find(f"{{{W_NS}}}b") is not None and bold_rpr is None:
            bold_rpr = deepcopy(rpr)
        if rpr.find(f"{{{W_NS}}}b") is None and nonbold_rpr is None:
            nonbold_rpr = deepcopy(rpr)
    return bold_rpr, nonbold_rpr


def _append_caption_run(p, text: str, rpr) -> None:
    if not text:
        return
    r_el = OxmlElement("w:r")
    if rpr is not None:
        r_el.append(deepcopy(rpr))
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_el.append(t_el)
    p._element.append(r_el)


def set_caption(doc: Document, caption_text: str, bold_rpr, nonbold_rpr) -> None:
    bold_prefix = "Table XXX."
    for p in doc.paragraphs:
        if p.text.strip():
            for child in list(p._element):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag != "pPr":
                    p._element.remove(child)
            if caption_text.startswith(bold_prefix):
                _append_caption_run(p, bold_prefix, bold_rpr)
                _append_caption_run(p, caption_text[len(bold_prefix):], nonbold_rpr)
            else:
                _append_caption_run(p, caption_text, nonbold_rpr)
            break


def clear_cell(cell, text: str, italic: bool = False, indent: bool = False) -> None:
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]

    # Remove all runs from all paragraphs in this cell, then create one run.
    for para in cell.paragraphs:
        for run in list(para.runs):
            run._element.getparent().remove(run._element)

    # Apply paragraph indent (w:firstLine 142) to match template category rows.
    if indent:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLine"), "142")
    else:
        # Remove any existing indent so header rows have none.
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                pPr.remove(ind)

    run = p.add_run(text)
    run.bold = False
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def insert_table_row_after(table, row_idx: int) -> None:
    """Insert a duplicate of row_idx immediately after it in the table XML."""
    ref_tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)


def delete_table_row(table, row_idx: int) -> None:
    """Remove a row from the table by index."""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)

def fill_row(table, row_idx: int, u_or: str, u_ci: str, m_or: str, m_ci: str) -> None:
    row = table.rows[row_idx]
    clear_cell(row.cells[1], u_or)
    clear_cell(row.cells[2], u_ci)
    clear_cell(row.cells[4], m_or)
    clear_cell(row.cells[5], m_ci)


def get_vals(data: dict[tuple[str, str], dict[str, str]], variable: str, level: str):
    row = data.get((variable, level), {})
    if not row:
        return "", "", "", ""
    u_or = format_or(row.get("univ_or", ""))
    u_ci = format_ci(row.get("univ_95ci", "")) if row.get("univ_95ci") else ""
    m_or = format_or(row.get("multiv_or", ""))
    m_ci = format_ci(row.get("multiv_95ci", "")) if row.get("multiv_95ci") else ""
    if u_or == "Ref":
        u_ci = ""
    if m_or == "Ref":
        m_ci = ""
    return u_or, u_ci, m_or, m_ci


def build_doc(
    template_path: Path,
    out_path: Path,
    data: dict[tuple[str, str], dict[str, str]],
    include_top_cited_row: bool,
    pub_label: str,
    caption: str,
) -> None:
    template_doc = Document(str(template_path))
    bold_rpr, nonbold_rpr = get_caption_rprs(template_doc)
    doc = Document(str(template_path))
    table = doc.tables[0]

    is_tertile = "tertile" in pub_label.lower()

    # For tertile mode: insert 3 data rows directly after the template's single
    # pub-volume row (18), then use row 17 (the empty spacer) as section header.
    # All original rows > 18 shift down by 3.
    PUB_ROW = 18
    if is_tertile:
        for _ in range(3):
            insert_table_row_after(table, PUB_ROW)
        pub_shift = 3
    else:
        pub_shift = 0

    def S(r: int) -> int:
        """Shift original template row index for rows inserted after PUB_ROW."""
        return r + pub_shift if r > PUB_ROW else r

    mappings = [
        (4, "Gender", "Men"),
        (9, "Age career (year of first publication)", "1992-2001"),
        (10, "Age career (year of first publication)", "2002-2011"),
        (11, "Age career (year of first publication)", ">=2012"),
        (15, "Income level", "All other income levels"),
        (16, "Income level", "Unknown"),
        (S(25), "Scientific field", "Agriculture, Fisheries & Forestry"),
        (S(26), "Scientific field", "Biology"),
        (S(27), "Scientific field", "Biomedical Research"),
        (S(28), "Scientific field", "Built Environment & Design"),
        (S(29), "Scientific field", "Chemistry"),
        (S(31), "Scientific field", "Communication & Textual Studies"),
        (S(32), "Scientific field", "Earth & Environmental Sciences"),
        (S(33), "Scientific field", "Economics & Business"),
        (S(34), "Scientific field", "Enabling & Strategic Technologies"),
        (S(35), "Scientific field", "Engineering"),
        (S(36), "Scientific field", "Historical Studies"),
        (S(37), "Scientific field", "Information & Communication Technologies"),
        (S(38), "Scientific field", "Mathematics & Statistics"),
        (S(39), "Scientific field", "Philosophy & Theology"),
        (S(40), "Scientific field", "Physics & Astronomy"),
        (S(41), "Scientific field", "Psychology & Cognitive Sciences"),
        (S(42), "Scientific field", "Public Health & Health Services"),
        (S(43), "Scientific field", "Social Sciences"),
        (S(44), "Scientific field", "Visual & Performing Arts"),
        (S(47), "Interaction", "Men in youngest cohort>=2012"),
    ]
    if include_top_cited_row:
        mappings.insert(7, (S(22), "Top-cited status", "Yes"))

    for row_idx, var, lvl in mappings:
        u_or, u_ci, m_or, m_ci = get_vals(data, var, lvl)
        fill_row(table, row_idx, u_or, u_ci, m_or, m_ci)

    # Publication volume rows
    if is_tertile:
        # Row 17 stays as empty spacer (not touched).
        # Row 18 becomes section header only (no data values) — italic.
        clear_cell(table.rows[PUB_ROW].cells[0], pub_label, italic=True)
        # Rows 19, 20, 21 are the 3 inserted tertile data rows — indented.
        for i, lvl in enumerate(["Tertile 1 (lowest)", "Tertile 2", "Tertile 3 (highest)"]):
            r = PUB_ROW + 1 + i
            clear_cell(table.rows[r].cells[0], lvl, indent=True)
            u_or, u_ci, m_or, m_ci = get_vals(data, pub_label, lvl)
            fill_row(table, r, u_or, u_ci, m_or, m_ci)
    else:
        u_or, u_ci, m_or, m_ci = get_vals(data, pub_label, "Per +1 unit")
        fill_row(table, PUB_ROW, u_or, u_ci, m_or, m_ci)
        clear_cell(table.rows[PUB_ROW].cells[0], pub_label)

    # Remove entire Top-cited status section when not applicable (top-cited cohort).
    # Delete in reverse order to keep indices stable: spacer after, Yes, No/Ref, header.
    if not include_top_cited_row:
        for blank_row in sorted([S(20), S(21), S(22), S(23)], reverse=True):
            delete_table_row(table, blank_row)

    set_caption(doc, caption, bold_rpr, nonbold_rpr)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def infer_pub_label(all_data: dict[tuple[str, str], dict[str, str]]) -> str:
    pub_keys = [k for k in all_data if k[0].startswith("Publication volume")]
    if not pub_keys:
        return "Publication volume (per paper)"
    return pub_keys[0][0]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--csv-base", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--suffix", default="")
    args = parser.parse_args()

    template = Path(args.template)
    base = Path(args.csv_base)
    out_dir = Path(args.output_dir)

    all_data = read_or_rows(str(base / "all_authors_or_table.csv"))
    top_data = read_or_rows(str(base / "top_cited_or_table.csv"))
    pub_label = infer_pub_label(all_data)

    if "tertile" in pub_label.lower():
        pub_text = "Publication volume is categorised into tertiles (T1 lowest, T3 highest)."
    elif "paper-year" in pub_label.lower():
        pub_text = "Publication volume is modelled as a continuous variable (OR per additional active publication year)."
    else:
        pub_text = "Publication volume is modelled as a continuous variable (OR per additional paper)."

    suffix = f"_{args.suffix}" if args.suffix else ""

    pub_ref = "; Tertile 1 (lowest)" if "tertile" in pub_label.lower() else ""

    all_caption = (
        "Table XXX. Univariable and multivariable odds ratios (OR) and 95% confidence intervals (CI) "
        "from logistic regression models estimating the probability of having >=1 retraction, among all authors. "
        f"{pub_text} Reference categories: Women; <1992; High income level; No top-cited status{pub_ref}; Clinical Medicine."
    )
    top_caption = (
        "Table XXX. Univariable and multivariable odds ratios (OR) and 95% confidence intervals (CI) "
        "from logistic regression models estimating the probability of having >=1 retraction, among top-cited authors. "
        f"{pub_text} Reference categories: Women; <1992; High income level{pub_ref}; Clinical Medicine."
    )

    build_doc(
        template,
        out_dir / f"Table_all_authors_OR{suffix}.docx",
        all_data,
        include_top_cited_row=True,
        pub_label=pub_label,
        caption=all_caption,
    )
    build_doc(
        template,
        out_dir / f"Table_top_cited_authors_OR{suffix}.docx",
        top_data,
        include_top_cited_row=False,
        pub_label=pub_label,
        caption=top_caption,
    )

    print("Generated:", out_dir / f"Table_all_authors_OR{suffix}.docx")
    print("Generated:", out_dir / f"Table_top_cited_authors_OR{suffix}.docx")


if __name__ == "__main__":
    main()
