#!/usr/bin/env python3
"""Build field × male interaction DOCX tables from interaction OR CSV outputs.

Usage example:
    .venv/bin/python projects/2026_INTERNAL_retraction/scripts/build_field_interaction_docx.py \
        --csv-base projects/2026_INTERNAL_retraction/output/local_or_tables/papers_log10 \
        --output-dir projects/2026_INTERNAL_retraction/output \
        --suffix papers_log10_local_20260320
"""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


def _format_ci(ci_str: str) -> str:
    ci_str = ci_str.strip()
    if ci_str.startswith("[") and ci_str.endswith("]"):
        inner = ci_str[1:-1]
        parts = inner.split(",")
        if len(parts) == 2:
            return f"{parts[0].strip()}-{parts[1].strip()}"
    return ci_str


def _set_cell_text(cell, text: str, bold: bool = False, italic: bool = False, indent: bool = False) -> None:
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    if indent:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._element.insert(0, pPr)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLine"), "142")

    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def build_interaction_doc(csv_path: Path, caption: str, cohort_label: str) -> Document:
    doc = Document()

    # Remove default blank paragraph that Document() adds
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)

    # Caption paragraph
    cap_para = doc.add_paragraph()
    bold_prefix = "Table XXX."
    r_bold = cap_para.add_run(bold_prefix)
    r_bold.bold = True
    r_bold.font.name = "Times New Roman"
    r_bold.font.size = Pt(11)
    r_rest = cap_para.add_run(caption[len(bold_prefix):])
    r_rest.font.name = "Times New Roman"
    r_rest.font.size = Pt(11)

    # Read CSV
    rows = []
    with open(csv_path, encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            rows.append(row)

    # Table: Scientific field | Interaction OR | 95% CI
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    # Set column widths
    for i, width_inches in enumerate([3.2, 1.2, 1.5]):
        for row in table.rows:
            row.cells[i].width = Inches(width_inches)

    # Header row
    hdr = table.rows[0]
    _set_cell_text(hdr.cells[0], "Scientific field", bold=True)
    _set_cell_text(hdr.cells[1], "Interaction OR", bold=True)
    _set_cell_text(hdr.cells[2], "95% CI", bold=True)

    # Data rows
    for i, row in enumerate(rows):
        tr = table.rows[i + 1]
        field = row["field"]
        or_val = row.get("interaction_or", "")
        ci_val = _format_ci(row.get("interaction_95ci", ""))

        is_ref = or_val == "Ref"
        is_clinical = field == "Clinical Medicine"

        if is_clinical:
            _set_cell_text(tr.cells[0], field, italic=False)
            _set_cell_text(tr.cells[1], "Ref")
            _set_cell_text(tr.cells[2], "")
        else:
            _set_cell_text(tr.cells[0], field, indent=True)
            _set_cell_text(tr.cells[1], or_val)
            _set_cell_text(tr.cells[2], ci_val if not is_ref else "")

    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-base", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--suffix", default="")
    args = parser.parse_args()

    base = Path(args.csv_base)
    out_dir = Path(args.output_dir)
    suffix = f"_{args.suffix}" if args.suffix else ""

    all_csv = base / "all_authors_field_male_interactions.csv"
    top_csv = base / "top_cited_field_male_interactions.csv"

    all_caption = (
        "Table XXX. Interaction odds ratios (OR) and 95% confidence intervals (CI) for male × "
        "scientific field terms from multivariable logistic regression, among all authors. "
        "An OR > 1 indicates a relatively higher retraction risk among men in that field "
        "compared to Clinical Medicine (reference for the interaction). "
        "All main covariates are included in the model."
    )
    top_caption = (
        "Table XXX. Interaction odds ratios (OR) and 95% confidence intervals (CI) for male × "
        "scientific field terms from multivariable logistic regression, among top-cited authors. "
        "An OR > 1 indicates a relatively higher retraction risk among men in that field "
        "compared to Clinical Medicine (reference for the interaction). "
        "All main covariates are included in the model."
    )

    out_dir.mkdir(parents=True, exist_ok=True)

    all_doc = build_interaction_doc(all_csv, all_caption, "all authors")
    top_doc = build_interaction_doc(top_csv, top_caption, "top-cited authors")

    all_out = out_dir / f"Table_all_authors_OR{suffix}_field_male_interactions.docx"
    top_out = out_dir / f"Table_top_cited_authors_OR{suffix}_field_male_interactions.docx"

    all_doc.save(str(all_out))
    top_doc.save(str(top_out))

    print(f"Generated: {all_out}")
    print(f"Generated: {top_out}")


if __name__ == "__main__":
    main()
