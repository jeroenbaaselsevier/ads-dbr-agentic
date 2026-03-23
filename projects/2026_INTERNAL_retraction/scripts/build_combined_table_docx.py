#!/usr/bin/env python3
"""Build combined 5-column OR table DOCX from combined_no_int.csv / combined_full_int.csv.

Combined CSV columns: variable, level, all_univ, all_multiv, top_univ, top_multiv
Each OR+CI is already merged as "1.23 (1.10-1.38)" or "Ref" or "".

Table layout (6 content columns + 1 label column = 7 cols total):
  Col 0: Variable / level label
  Col 1: All authors — Univariable OR (95% CI)
  Col 2: All authors — Multivariable OR (95% CI)
  Col 3: [spacer / divider]
  Col 4: Top-cited authors — Univariable OR (95% CI)
  Col 5: Top-cited authors — Multivariable OR (95% CI)

Usage:
  python projects/2026_INTERNAL_retraction/scripts/build_combined_table_docx.py \
    --csv combined_no_int.csv \
    --output Table_combined_papers_log10_no_int.docx \
    --pub-label "Publication volume (log10 papers)" \
    --caption "Table XXX. ..."
"""
from __future__ import annotations

import argparse
import csv
import glob
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


# ---------------------------------------------------------------------------
# Page settings
# ---------------------------------------------------------------------------

def _set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    w = max(section.page_width, section.page_height)
    h = min(section.page_width, section.page_height)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = w
    section.page_height = h
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


# ---------------------------------------------------------------------------
# Table border / layout helpers
# ---------------------------------------------------------------------------

def _set_table_borders(table) -> None:
    """Top and bottom outer borders only; no vertical or interior horizontal lines."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    tblPr.append(tblBorders)
    for name, show in [
        ("top", True), ("left", False), ("bottom", True), ("right", False),
        ("insideH", False), ("insideV", False),
    ]:
        el = OxmlElement(f"w:{name}")
        if show:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        tblBorders.append(el)


def _set_table_autofit(table) -> None:
    """Set table to full page width."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


# Column widths in twips for a 9" content area (landscape, 1" margins each side):
# Col 0: label 3.1", Cols 1-2: 1.4" each, Col 3: spacer 0.3", Cols 4-5: 1.4" each
# Total: 3.1 + 1.4*4 + 0.3 = 9.0" = 12960 twips
# Column widths in twips for a 9\" content area (landscape, 1\" margins each side):
# Col 0: label 2.4\", Cols 1-2: counts 0.5\" each, Cols 3-4: OR 1.05\" each,
# Col 5: spacer 0.2\", Cols 6-7: counts 0.5\" each, Cols 8-9: OR 1.05\" each
# Total: 2.4 + 0.5*2 + 1.05*2 + 0.2 + 0.5*2 + 1.05*2 = 9.0\" = 12960 twips
_COL_WIDTHS_TWIPS = [3456, 720, 720, 1512, 1512, 288, 720, 720, 1512, 1512]
_N_COLS = 10


def _set_col_widths(table) -> None:
    """Set column widths in twips via tblGrid and per-cell tcW (handles merged cells)."""
    # Set tblGrid widths
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for col, width in zip(tblGrid.findall(qn("w:gridCol")), _COL_WIDTHS_TWIPS):
            col.set(qn("w:w"), str(width))
    # Set per-cell tcW, summing spans for merged cells
    for tr in table._tbl.findall(qn("w:tr")):
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            gridSpan = tcPr.find(qn("w:gridSpan"))
            span = int(gridSpan.get(qn("w:val"), 1)) if gridSpan is not None else 1
            cell_width = sum(_COL_WIDTHS_TWIPS[col_idx:col_idx + span])
            existing = tcPr.find(qn("w:tcW"))
            if existing is not None:
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(cell_width))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            col_idx += span


def _add_cell_bottom_border(cell) -> None:
    """Add a single bottom border to a cell (used for header row underline)."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    tcPr.append(tcBorders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    tcBorders.append(bottom)


def _set_row_repeat_header(row) -> None:
    """Mark row to repeat as table header on every page."""
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        row._tr.insert(0, trPr)
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


# ---------------------------------------------------------------------------
# Cell text helper
# ---------------------------------------------------------------------------

def _set_cell(cell, text: str, bold: bool = False, italic: bool = False,
              indent: bool = False, center: bool = False) -> None:
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    pPr = para._element.find(qn("w:pPr"))
    if indent:
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._element.insert(0, pPr)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "152")
    else:
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                pPr.remove(ind)

    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Row classification
# ---------------------------------------------------------------------------

SECTION_HEADER_VARIABLES = {
    "Gender",
    "Age career (year of first publication)",
    "Income level",
    "Top-cited status",
    "Scientific field",
    "Interaction",
}

PUB_VOLUME_PREFIXES = ("Publication volume",)


def _is_degenerate(val: str) -> bool:
    """True if a cell value represents a numerically meaningless estimate."""
    v = (val or "").strip()
    if not v or v == "Ref":
        return False
    return (
        "1.00 (1.00-1.00)" in v
        or ">999" in v
        or v.startswith("<0.01 (<0.01->")
    )


def _sanitize(val: str) -> str:
    """Replace degenerate cell values with an em-dash for display."""
    return "\u2014" if _is_degenerate(val) else (val or "")


def _row_should_drop(row: dict) -> bool:
    """Drop only rows where ALL non-empty cells are exactly '1.00 (1.00-1.00)',
    which indicates the dummy was never 1 in the data (field not in dataset).
    """
    if row.get("_type") != "data":
        return False
    cells = [row.get(c, "") for c in ("all_univ", "all_multiv", "top_univ", "top_multiv")]
    if any(str(c).strip() == "Ref" for c in cells):
        return False
    non_empty = [c for c in cells if (c or "").strip()]
    if not non_empty:
        return False
    return all(str(c).strip() == "1.00 (1.00-1.00)" for c in non_empty)


def _is_ref(all_univ: str) -> bool:
    return all_univ.strip() == "Ref"


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def read_combined_csv(path: Path) -> list[dict]:
    rows = []
    with open(path, encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            rows.append(row)
    return rows


_PARQUET_MAPPINGS = [
    ("Gender", "gender_clean", {
        "female": "Women", "male": "Men", "unknown": "Unknown",
    }),
    ("Age career (year of first publication)", "career_age_clean", None),
    ("Income level", "income_clean", {
        "High Income": "High income level",
        "All Other Income Levels": "All other income levels",
        "unknown": "Unknown",
    }),
    ("Scientific field", "field_clean", None),
    ("Top-cited status", "top_cited_yes", {0: "No", 1: "Yes"}),
]


def _load_parquet_counts(parquet_dir: Path) -> dict:
    """Return {(variable, level): (all_n, all_n_ret, top_n, top_n_ret)} from grouped parquet."""
    import pandas as pd

    files = sorted(glob.glob(str(parquet_dir / "*.parquet")))
    if not files:
        return {}
    df = pd.concat([pd.read_parquet(f) for f in files], ignore_index=True)

    retracted = df["label"] == 1
    top_cited = df["top_cited_yes"] == 1

    result = {}

    for variable, col, val_map in _PARQUET_MAPPINGS:
        if col not in df.columns:
            continue
        for pval in df[col].unique():
            mask = df[col] == pval
            all_n     = int(df.loc[mask,              "n_obs"].sum())
            all_n_ret = int(df.loc[mask & retracted,  "n_obs"].sum())
            top_n     = int(df.loc[mask & top_cited,  "n_obs"].sum())
            top_n_ret = int(df.loc[mask & retracted & top_cited, "n_obs"].sum())
            if val_map is not None:
                display = val_map.get(pval)
                if display is None:
                    try:
                        display = val_map.get(int(pval))
                    except (TypeError, ValueError):
                        display = None
            else:
                display = str(pval)
            if display:
                result[(variable, display)] = (all_n, all_n_ret, top_n, top_n_ret)

    # Publication volume rows get total counts
    result[("__pub_total__",)] = (
        int(df["n_obs"].sum()),
        int(df.loc[retracted,            "n_obs"].sum()),
        int(df.loc[top_cited,            "n_obs"].sum()),
        int(df.loc[retracted & top_cited, "n_obs"].sum()),
    )

    # Interaction rows: each level is a conjunction of two parquet column filters.
    # For the "all authors" side: count everyone matching the filters.
    # For the "top-cited" side: additionally require top_cited_yes == 1.
    _INT_FILTERS: list[tuple[str, list[tuple[str, object]]]] = [
        ("Men in youngest cohort (>=2012)", [("gender_clean", "male"), ("career_age_clean", ">=2012")]),
        ("Men in all other income levels",  [("gender_clean", "male"), ("income_clean", "All Other Income Levels")]),
        ("Men in unknown income",           [("gender_clean", "male"), ("income_clean", "unknown")]),
        ("Men × publication volume",        [("gender_clean", "male")]),
        ("Men × top-cited",                 [("gender_clean", "male"), ("top_cited_yes", 1.0)]),
    ]
    # Add Men × {field} for every field present in the parquet
    for fv in df["field_clean"].dropna().unique():
        _INT_FILTERS.append((f"Men \u00d7 {fv}", [("gender_clean", "male"), ("field_clean", fv)]))

    for level_str, filters in _INT_FILTERS:
        mask = True
        for col, val in filters:
            mask = mask & (df[col] == val)
        # For top-cited side, also require top_cited unless already filtered
        already_top = any(col == "top_cited_yes" for col, _ in filters)
        top_mask = mask if already_top else (mask & top_cited)
        result[("Interaction", level_str)] = (
            int(df.loc[mask,                    "n_obs"].sum()),
            int(df.loc[mask & retracted,        "n_obs"].sum()),
            int(df.loc[top_mask,                "n_obs"].sum()),
            int(df.loc[top_mask & retracted,    "n_obs"].sum()),
        )

    return result


def build_table_doc(rows: list[dict], caption: str, pub_label: str, include_top_cited_section: bool) -> Document:
    doc = Document()
    _set_landscape(doc)

    # Remove default blank paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Caption paragraph
    cap_para = doc.add_paragraph()
    rb = cap_para.add_run("Table XXX.")
    rb.bold = True
    rb.font.name = "Times New Roman"
    rb.font.size = Pt(10)
    rest = caption[len("Table XXX."):]
    rr = cap_para.add_run(rest)
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(10)

    # Determine rows to display, injecting section-header rows
    display_rows: list[dict] = []
    prev_variable = None
    for row in rows:
        variable = row["variable"]

        # Skip Top-cited section rows when excluded
        if not include_top_cited_section and variable == "Top-cited status":
            continue

        # Inject section-header synthetic row when variable changes
        if variable != prev_variable:
            if variable in SECTION_HEADER_VARIABLES or any(variable.startswith(p) for p in PUB_VOLUME_PREFIXES):
                display_rows.append({"_type": "header", "text": variable})
            prev_variable = variable

        display_rows.append({"_type": "data", **row})

    # Drop fully-degenerate rows (no-data dummies, completely failed estimates)
    display_rows = [r for r in display_rows if not _row_should_drop(r)]

    # Build table: 2 header rows + data rows
    n_rows = 2 + len(display_rows)
    table = doc.add_table(rows=n_rows, cols=_N_COLS)
    _set_table_autofit(table)
    _set_table_borders(table)

    # --- Header row 0: group labels spanning 4 cols each ---
    hdr0 = table.rows[0]
    _set_row_repeat_header(hdr0)
    _set_cell(hdr0.cells[0], "", bold=True)
    hdr0.cells[1].merge(hdr0.cells[4])   # cols 1-4: All authors
    _set_cell(hdr0.cells[1], "All authors", bold=True, center=True)
    _set_cell(hdr0.cells[5], "", bold=True)  # spacer
    hdr0.cells[6].merge(hdr0.cells[9])   # cols 6-9: Top-cited authors
    _set_cell(hdr0.cells[6], "Top-cited authors", bold=True, center=True)

    # --- Header row 1: sub-labels with bottom border ---
    hdr1 = table.rows[1]
    _set_row_repeat_header(hdr1)
    _set_cell(hdr1.cells[0], "", bold=True)
    _set_cell(hdr1.cells[1], "n", bold=True, center=True)
    _set_cell(hdr1.cells[2], "n retracted", bold=True, center=True)
    _set_cell(hdr1.cells[3], "Univariable OR (95% CI)", bold=True, center=True)
    _set_cell(hdr1.cells[4], "Multivariable OR (95% CI)", bold=True, center=True)
    _set_cell(hdr1.cells[5], "", bold=True)
    _set_cell(hdr1.cells[6], "n", bold=True, center=True)
    _set_cell(hdr1.cells[7], "n retracted", bold=True, center=True)
    _set_cell(hdr1.cells[8], "Univariable OR (95% CI)", bold=True, center=True)
    _set_cell(hdr1.cells[9], "Multivariable OR (95% CI)", bold=True, center=True)
    for cell in hdr1.cells:
        _add_cell_bottom_border(cell)

    _set_col_widths(table)  # apply after merges

    # --- Data rows (no shading) ---
    for i, drow in enumerate(display_rows):
        tr = table.rows[i + 2]  # offset by 2 header rows
        rtype = drow.get("_type", "data")

        if rtype == "header":
            _set_cell(tr.cells[0], drow["text"], italic=True)
            for ci in range(1, _N_COLS):
                _set_cell(tr.cells[ci], "")

        else:
            level = drow.get("level", "")
            all_univ   = _sanitize(drow.get("all_univ", ""))
            all_multiv = _sanitize(drow.get("all_multiv", ""))
            top_univ   = _sanitize(drow.get("top_univ", ""))
            top_multiv = _sanitize(drow.get("top_multiv", ""))

            def _fmt_count(val) -> str:
                try:
                    v = int(float(val))
                    return f"{v:,}"
                except (TypeError, ValueError):
                    return ""

            all_n     = _fmt_count(drow.get("all_n", ""))
            all_n_ret = _fmt_count(drow.get("all_n_ret", ""))
            top_n     = _fmt_count(drow.get("top_n", ""))
            top_n_ret = _fmt_count(drow.get("top_n_ret", ""))

            _set_cell(tr.cells[0], level, indent=True)
            _set_cell(tr.cells[1], all_n,     center=True)
            _set_cell(tr.cells[2], all_n_ret, center=True)
            _set_cell(tr.cells[3], all_univ,   center=True)
            _set_cell(tr.cells[4], all_multiv, center=True)
            _set_cell(tr.cells[5], "")  # spacer
            _set_cell(tr.cells[6], top_n,     center=True)
            _set_cell(tr.cells[7], top_n_ret, center=True)
            _set_cell(tr.cells[8], top_univ,   center=True)
            _set_cell(tr.cells[9], top_multiv, center=True)

    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv", required=True, help="Path to combined_no_int.csv or combined_full_int.csv")
    parser.add_argument("--parquet", default=None, help="Directory of grouped parquet files for count columns")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--pub-label", default="Publication volume", help="Label for pub volume section header")
    parser.add_argument("--caption", default="", help="Full caption text starting with 'Table XXX.'")
    parser.add_argument("--include-top-cited-section", action="store_true", default=True,
                        help="Include Top-cited status rows (default: True for all-authors tables)")
    parser.add_argument("--no-top-cited-section", dest="include_top_cited_section", action="store_false")
    args = parser.parse_args()

    csv_path = Path(args.csv)
    out_path = Path(args.output)

    rows = read_combined_csv(csv_path)

    if args.parquet:
        pq_counts = _load_parquet_counts(Path(args.parquet))
        for row in rows:
            variable = row["variable"]
            level    = row["level"]
            if any(variable.startswith(p) for p in PUB_VOLUME_PREFIXES):
                c = pq_counts.get(("__pub_total__",))
            else:
                c = pq_counts.get((variable, level))
            if c:
                row["all_n"], row["all_n_ret"], row["top_n"], row["top_n_ret"] = c

    doc = build_table_doc(rows, args.caption, args.pub_label, args.include_top_cited_section)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
